import base64
from io import BytesIO
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from models.portfolio import Portfolio
from models.user_profile import UserProfile
from models.resume import TailoredResume
from models.career_description import CareerDescription


class DocumentGenerator:
    """Clean, professional resume/career document generator."""

    PRIMARY = RGBColor(0x1A, 0x56, 0xDB)   # 깔끔한 파랑
    DARK = RGBColor(0x1F, 0x1F, 0x1F)
    GRAY = RGBColor(0x5A, 0x5A, 0x5A)
    LIGHT_GRAY = RGBColor(0x8A, 0x8A, 0x8A)

    def _set_cell_border(self, cell, **kwargs):
        """Set cell borders. Usage: _set_cell_border(cell, bottom={"sz": 4, "color": "000000"})"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.first_child_found_in(qn('w:tcBorders'))
        if tcBorders is None:
            tcBorders = tcPr.makeelement(qn('w:tcBorders'), {})
            tcPr.append(tcBorders)
        for edge, attrs in kwargs.items():
            element = tcBorders.makeelement(qn(f'w:{edge}'), {
                qn('w:val'): 'single',
                qn('w:sz'): str(attrs.get('sz', 4)),
                qn('w:space'): '0',
                qn('w:color'): attrs.get('color', '000000'),
            })
            tcBorders.append(element)

    def _add_thin_line(self, doc, color='CCCCCC'):
        """Add a subtle thin divider line."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn('w:pBdr'), {})
        bottom = pBdr.makeelement(qn('w:bottom'), {
            qn('w:val'): 'single',
            qn('w:sz'): '4',
            qn('w:space'): '1',
            qn('w:color'): color,
        })
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _section_heading(self, doc, text: str):
        """Section heading with primary color and bottom border."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = self.PRIMARY
        # subtle line
        self._add_thin_line(doc, '1A56DB')

    def _entry_title(self, doc, bold_text: str, light_text: str = ""):
        """Entry title with bold name and light metadata."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(bold_text)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = self.DARK
        if light_text:
            run = p.add_run(f"   {light_text}")
            run.font.size = Pt(9)
            run.font.color.rgb = self.LIGHT_GRAY

    def _body_text(self, doc, text: str):
        """Regular body text."""
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(9.5)
            run.font.color.rgb = self.GRAY

    def _bullet(self, doc, text: str):
        """Bullet point item."""
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.8)
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.color.rgb = self.GRAY

    def _tag_line(self, doc, text: str):
        """Small gray tag text (tech stack, etc.)."""
        p = doc.add_paragraph(text)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = self.LIGHT_GRAY

    def _decode_photo(self, photo_url: str) -> BytesIO | None:
        if not photo_url or not photo_url.startswith("data:"):
            return None
        try:
            _, b64_data = photo_url.split(",", 1)
            image_bytes = base64.b64decode(b64_data)
            return BytesIO(image_bytes)
        except Exception:
            return None

    def _add_photo_header(self, doc, profile: UserProfile, title_text: str, subtitle_text: str = ""):
        photo_buf = self._decode_photo(profile.photo_url) if hasattr(profile, 'photo_url') else None

        if photo_buf:
            table = doc.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.autofit = False

            tbl = table._tbl
            tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
            borders = tblPr.makeelement(qn('w:tblBorders'), {})
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border_elem = borders.makeelement(qn(f'w:{border_name}'), {
                    qn('w:val'): 'none', qn('w:sz'): '0', qn('w:space'): '0', qn('w:color'): 'auto'
                })
                borders.append(border_elem)
            tblPr.append(borders)

            photo_cell = table.cell(0, 0)
            photo_cell.width = Cm(3)
            photo_para = photo_cell.paragraphs[0]
            photo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = photo_para.add_run()
            run.add_picture(photo_buf, width=Cm(2.5))

            info_cell = table.cell(0, 1)
            p = info_cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            run = p.add_run(title_text)
            run.bold = True
            run.font.size = Pt(22)
            run.font.color.rgb = self.DARK

            contact = self._build_contact(profile)
            if contact:
                p2 = info_cell.add_paragraph(" · ".join(contact))
                p2.paragraph_format.space_before = Pt(4)
                p2.paragraph_format.space_after = Pt(2)
                for r in p2.runs:
                    r.font.size = Pt(9)
                    r.font.color.rgb = self.LIGHT_GRAY

            if subtitle_text:
                p3 = info_cell.add_paragraph(subtitle_text)
                p3.paragraph_format.space_after = Pt(2)
                for r in p3.runs:
                    r.font.size = Pt(10)
                    r.font.color.rgb = self.GRAY
            return True
        return False

    def _build_contact(self, profile):
        contact = []
        if profile.email:
            contact.append(profile.email)
        if profile.phone:
            contact.append(profile.phone)
        if profile.github:
            contact.append(profile.github)
        if profile.blog:
            contact.append(profile.blog)
        return contact

    def _add_text_header(self, doc, profile, title_text, subtitle_text=""):
        """Add text-only header (no photo)."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(title_text)
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = self.DARK

        contact = self._build_contact(profile)
        if contact:
            p = doc.add_paragraph(" · ".join(contact))
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = self.LIGHT_GRAY

        if subtitle_text:
            p = doc.add_paragraph(subtitle_text)
            p.paragraph_format.space_after = Pt(4)
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = self.GRAY

        self._add_thin_line(doc, '333333')

    def generate_docx(
        self,
        profile: UserProfile,
        tailored_resume: TailoredResume,
        portfolios: list[Portfolio],
    ) -> BytesIO:
        doc = Document()

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.line_spacing = 1.25

        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # === Header (Name + Contact) ===
        has_photo = self._add_photo_header(doc, profile, profile.name)
        if not has_photo:
            self._add_text_header(doc, profile, profile.name)

        # === Summary ===
        if tailored_resume.summary:
            self._section_heading(doc, "Summary")
            self._body_text(doc, tailored_resume.summary)

        # === Experience ===
        portfolio_map = {p.id: p for p in portfolios}
        valid_entries = [e for e in tailored_resume.entries if portfolio_map.get(e.portfolio_id)]
        has_work = profile.work_experience and any(w.company for w in profile.work_experience)

        if has_work or valid_entries:
            self._section_heading(doc, "Experience")

            if has_work:
                for work in profile.work_experience:
                    if not work.company:
                        continue
                    period = f"{work.start_date} ~ {'재직중' if work.is_current else work.end_date}"
                    self._entry_title(doc, f"{work.company} · {work.position}", period)
                    if work.team:
                        self._body_text(doc, work.team)
                    if work.description:
                        self._bullet(doc, work.description)
                    for proj in work.projects:
                        if proj.name:
                            self._bullet(doc, f"{proj.name}: {proj.description}")

            if valid_entries:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run("주요 프로젝트")
                run.bold = True
                run.font.size = Pt(10.5)
                run.font.color.rgb = self.DARK

                for entry in valid_entries:
                    portfolio = portfolio_map[entry.portfolio_id]
                    self._entry_title(
                        doc,
                        portfolio.title,
                        f"{portfolio.role} · {portfolio.period}",
                    )
                    self._body_text(doc, entry.tailored_description)
                    for ach in entry.tailored_achievements:
                        self._bullet(doc, ach)
                    if portfolio.tech_stack:
                        self._tag_line(doc, f"Tech: {', '.join(portfolio.tech_stack)}")

        # === Skills ===
        all_skills: set[str] = set()
        for portfolio in portfolios:
            all_skills.update(portfolio.tech_stack)
        if all_skills:
            self._section_heading(doc, "Skills")
            p = doc.add_paragraph(", ".join(sorted(all_skills)))
            p.paragraph_format.space_after = Pt(4)
            for run in p.runs:
                run.font.size = Pt(9.5)
                run.font.color.rgb = self.GRAY

        # === Education ===
        if profile.education:
            self._section_heading(doc, "Education")
            for edu in profile.education:
                parts = []
                if edu.major:
                    parts.append(edu.major)
                if edu.degree:
                    parts.append(edu.degree)
                period = ""
                if edu.start_date or edu.end_date:
                    period = f"{edu.start_date} ~ {edu.end_date}"
                gpa_text = f"  GPA {edu.gpa}/{edu.gpa_scale}" if edu.gpa else ""
                self._entry_title(
                    doc,
                    edu.school,
                    f"{' · '.join(parts)}  {period}{gpa_text}",
                )

        # === Certifications ===
        if profile.certificates:
            has_valid = any(c.name for c in profile.certificates)
            if has_valid:
                self._section_heading(doc, "Certifications")
                for cert in profile.certificates:
                    if not cert.name:
                        continue
                    self._entry_title(doc, cert.name, f"{cert.issuer}  {cert.date}")

        # === Awards ===
        if profile.awards:
            has_valid = any(a.name for a in profile.awards)
            if has_valid:
                self._section_heading(doc, "Awards")
                for award in profile.awards:
                    if not award.name:
                        continue
                    self._entry_title(doc, award.name, f"{award.issuer}  {award.date}")
                    if award.description:
                        self._body_text(doc, award.description)

        # === Training ===
        if profile.trainings:
            has_valid = any(t.name for t in profile.trainings)
            if has_valid:
                self._section_heading(doc, "Training")
                for trn in profile.trainings:
                    if not trn.name:
                        continue
                    period = ""
                    if trn.start_date or trn.end_date:
                        period = f"{trn.start_date} ~ {trn.end_date}"
                    inst = f" / {trn.institution}" if trn.institution else ""
                    self._entry_title(doc, trn.name, f"{period}{inst}")
                    if trn.description:
                        self._body_text(doc, trn.description)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def generate_cover_letter_docx(
        self, answers: list[dict], profile: UserProfile | None = None
    ) -> BytesIO:
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.line_spacing = 1.4

        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3)
            section.right_margin = Cm(3)

        if profile and profile.name:
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title.add_run(f"{profile.name}")
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = self.DARK

            subtitle = doc.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.paragraph_format.space_after = Pt(12)
            run = subtitle.add_run("자기소개서")
            run.font.size = Pt(12)
            run.font.color.rgb = self.LIGHT_GRAY

            self._add_thin_line(doc)

        for i, qa in enumerate(answers, 1):
            q_para = doc.add_paragraph()
            q_para.paragraph_format.space_before = Pt(14)
            q_para.paragraph_format.space_after = Pt(6)
            run = q_para.add_run(f"Q{i}. {qa['question']}")
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = self.PRIMARY

            a_para = doc.add_paragraph(qa["answer"])
            a_para.paragraph_format.space_after = Pt(8)
            a_para.paragraph_format.line_spacing = 1.5
            for run in a_para.runs:
                run.font.size = Pt(10.5)
                run.font.color.rgb = self.GRAY

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def generate_career_desc_docx(
        self,
        profile: UserProfile,
        career_desc: CareerDescription,
        portfolios: list[Portfolio],
    ) -> BytesIO:
        doc = Document()

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.line_spacing = 1.25

        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # Title + Photo
        has_photo = self._add_photo_header(
            doc, profile,
            profile.name,
            f"경력기술서 · 지원 직무: {career_desc.target_role}"
        )

        if not has_photo:
            self._add_text_header(
                doc, profile,
                profile.name,
                f"경력기술서 · 지원 직무: {career_desc.target_role}"
            )

        # Summary
        if career_desc.summary:
            self._section_heading(doc, "경력 요약")
            self._body_text(doc, career_desc.summary)

        # Entries
        if career_desc.entries:
            self._section_heading(doc, "경력 상세")
            for entry in career_desc.entries:
                self._entry_title(doc, f"{entry.company} · {entry.position}", entry.period)

                if entry.description:
                    self._body_text(doc, entry.description)

                if entry.key_achievements:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(4)
                    run = p.add_run("핵심 성과")
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = self.DARK
                    for ach in entry.key_achievements:
                        self._bullet(doc, ach)

                if entry.relevant_projects:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(4)
                    run = p.add_run("관련 프로젝트")
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = self.DARK
                    for proj in entry.relevant_projects:
                        self._bullet(doc, proj)

                # spacing between entries
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf
